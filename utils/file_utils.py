import os
import io
import uuid
from datetime import datetime
from werkzeug.utils import secure_filename

# Existing imports
from db import db
from models import UploadedFile

# Document parsers
from docx import Document
from openpyxl import load_workbook
from PyPDF2 import PdfReader

# Optional Azure imports
try:
    from azure.storage.blob import BlobServiceClient
except ImportError:
    # If azure.storage.blob is not installed, handle gracefully
    BlobServiceClient = None

WORD_LIMIT = 50000

def process_uploaded_file(
    file=None,
    upload_folder=None,
    session_id=None,
    db_session=None,
    read=False,
    path=None,
    use_azure=False,
    blob_service_client=None,
    container_name=None
):
    """
    Handles file saving and processing.
    
    By default, files are stored locally (in `upload_folder`). 
    If `use_azure=True` and a valid `blob_service_client` is provided,
    files will be uploaded to Azure Blob Storage instead.

    :param file: File object from the request.
    :param upload_folder: Directory to save local uploads (if not using Azure).
    :param session_id: Current session ID.
    :param db_session: Database session.
    :param read: If True, read the file content (from `path`) but do not upload.
    :param path: Path to the file if reading.
    :param use_azure: Whether to upload to Azure instead of saving locally.
    :param blob_service_client: An instance of BlobServiceClient (if use_azure=True).
    :param container_name: The name of the Azure container (if use_azure=True).
    :return: Tuple (file_content, file_url, file_type, uploaded_file)
    """

    # 1) If 'read=True' and 'path' is provided, we just read the file content from disk
    #    and do not upload anywhere.
    print('In processing file', flush=True)
    if read and path:
        return read_file_content(path)

    # 2) If there's no file provided, we can't do anything
    print('No file', flush=True)
    if not file:
        print('Returning, no file', flush=True)
        return '', None, None, None

    print('Not returning, continuing to process files.', flush=True)
    # Generate secure filename
    filename = secure_filename(file.filename)
    unique_filename = f"{uuid.uuid4()}_{filename}"

    # --------------------------
    # A) Upload to Azure Storage
    # --------------------------
    if use_azure and blob_service_client and container_name:
        # Read file content into memory
        file_bytes = file.read()

        # Attempt uploading to Azure
        blob_client = blob_service_client.get_blob_client(
            container=container_name,
            blob=unique_filename
        )
        try:
            blob_client.upload_blob(file_bytes, overwrite=True)
            azure_blob_url = blob_client.url
            file_url = azure_blob_url
        except Exception as e:
            print("Error uploading to Azure Blob Storage:", e)
            return "Error uploading file.", None, None, None

        # Insert a record in the database with the Azure file info
        uploaded_file = UploadedFile(
            session_id=session_id,
            filename=unique_filename,
            original_filename=filename,
            file_url=file_url,
            file_type=file.content_type
        )
        db_session.session.add(uploaded_file)
        db_session.session.commit()

        # Extract the text for indexing/LLM usage
        file_content = extract_content_from_memory(
            file_bytes=file_bytes,
            content_type=file.content_type
        )
        file_type = file.content_type

    # -------------------------
    # B) Store File Locally
    # -------------------------
    else:
        # Construct local file path and save the file
        file_path = os.path.join(upload_folder, unique_filename)
        file.save(file_path)
        print(flush=True)

        # Log file save details
        try:
            file_size = os.path.getsize(file_path)
            print(f"Saved file at {file_path}, size: {file_size} bytes", flush=True)
        except Exception as e:
            print(f"Error accessing saved file: {e}", flush=True)
            return "Error accessing saved file.", None, None, None


        print(f'In process files util: {session_id}', flush=True)
        print(f'In process files util. file_path: {file_path}', flush=True)

        print(flush=True)

        # Insert record in DB with local file path
        uploaded_file = UploadedFile(
            session_id=session_id,
            filename=unique_filename,
            original_filename=filename,
            file_url=f"/uploads/{unique_filename}",
            file_type=file.content_type
        )
        print(f'Adding to db: {uploaded_file}', flush=True)
        db_session.session.add(uploaded_file)
        db_session.session.commit()
        print(f'Added to db: {uploaded_file}', flush=True)

        # Extract text from local file
        if file.content_type == 'application/pdf':
            file_content = extract_text_from_pdf(file_path)
        elif file.content_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]:
            file_content = extract_text_from_docx(file_path)
        elif file.content_type in [
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel'
        ]:
            file_content = extract_text_from_excel(file_path)
        else:
            # Attempt reading as text
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    file_content = f.read()
            except Exception as e:
                print("Error reading file:", e)
                file_content = "Error processing file."

        # Local URL for the stored file
        file_url = f"/uploads/{unique_filename}"
        file_type = file.content_type

    return file_content, file_url, file_type, uploaded_file


def read_file_content(path):
    """
    Reads an existing file from a local path, returning its text content.
    If you want to read from Azure for an existing file, you'll need to
    download to a temp file or memory first (not implemented here).
    """
    file_extension = os.path.splitext(path)[1].lower()
    if file_extension == '.pdf':
        return extract_text_from_pdf(path)
    elif file_extension in ['.docx', '.doc']:
        return extract_text_from_docx(path)
    elif file_extension in ['.xlsx', '.xls']:
        return extract_text_from_excel(path)
    else:
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print("Error reading file:", e)
            return "Error processing file."


# -----------------------
# In-memory extractions
# -----------------------
def extract_content_from_memory(file_bytes, content_type):
    """
    Reads the file content from an in-memory bytes object.
    """
    if not file_bytes:
        return ""

    if content_type == 'application/pdf':
        return extract_pdf_from_memory(file_bytes)
    elif content_type in [
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'
    ]:
        return extract_docx_from_memory(file_bytes)
    elif content_type in [
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'application/vnd.ms-excel'
    ]:
        return extract_excel_from_memory(file_bytes)
    else:
        # Assume it's plain text
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except Exception as e:
            print("Error decoding in-memory text file:", e)
            return "Error processing file."


def extract_pdf_from_memory(file_bytes):
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        file_content = ""
        for page in reader.pages:
            file_content += page.extract_text() or ""
        return truncate_content(file_content)
    except Exception as e:
        print("Error reading PDF from memory:", e)
        return "Error processing PDF file."


def extract_docx_from_memory(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))
        file_content = "\n".join([p.text for p in doc.paragraphs])
        return truncate_content(file_content)
    except Exception as e:
        print("Error reading DOCX from memory:", e)
        return "Error processing Word file."


def extract_excel_from_memory(file_bytes):
    try:
        wb = load_workbook(filename=io.BytesIO(file_bytes))
        sheet = wb.active
        file_content = ""
        for row in sheet.iter_rows(values_only=True):
            file_content += ' '.join(str(cell) for cell in row if cell is not None) + "\n"
        return truncate_content(file_content)
    except Exception as e:
        print("Error reading Excel from memory:", e)
        return "Error processing Excel file."


# -----------------------
# Local file extractions
# -----------------------
def extract_text_from_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        file_content = ""
        for page in reader.pages:
            file_content += page.extract_text() or ""
        return truncate_content(file_content)
    except Exception as e:
        print("Error reading PDF:", e)
        return "Error processing PDF file."


def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        file_content = "\n".join([p.text for p in doc.paragraphs])
        return truncate_content(file_content)
    except Exception as e:
        print("Error reading DOCX:", e)
        return "Error processing Word file."


def extract_text_from_excel(file_path):
    try:
        wb = load_workbook(file_path)
        sheet = wb.active
        file_content = ""
        for row in sheet.iter_rows(values_only=True):
            # Convert each cell to string if not None, then join
            file_content += ' '.join(str(cell) for cell in row if cell is not None) + "\n"
        return truncate_content(file_content)
    except Exception as e:
        print("Error reading Excel file:", e)
        return "Error processing Excel file."


def truncate_content(content):
    """
    Utility to ensure we don't exceed the WORD_LIMIT.
    """
    words = content.split()
    if len(words) > WORD_LIMIT:
        return ' '.join(words[:WORD_LIMIT]) + "\n\n[Text truncated after 50,000 words.]"
    return content
