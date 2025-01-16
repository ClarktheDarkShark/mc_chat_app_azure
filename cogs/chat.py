# cogs/chat.py
from flask import Blueprint, request, jsonify, send_from_directory, copy_current_request_context, session
import os
import gc
import json
import uuid
import openai
import tiktoken
from werkzeug.utils import secure_filename
from db import db
import traceback
from models import Conversation, Message, UploadedFile
from datetime import datetime
# Import the updated process_uploaded_file with Azure support
# from utils.file_utils import process_uploaded_file
from flask_socketio import rooms
from sqlalchemy.orm import joinedload

from cogs.orchestration_analysis import OrchestrationAnalysisCog
from utils.response_generation import generate_image, generate_chat_response
from .web_search import WebSearchCog
from .code_files import CodeFilesCog
from cogs.code_structure_visualizer import CodeStructureVisualizerCog  # New import
# from langgraph import LangGraph

import io
from models import UploadedFile

# Document parsers
from docx import Document
from openpyxl import load_workbook
from PyPDF2 import PdfReader


# Azure
from azure.identity import DefaultAzureCredential
from azure.keyvault.secrets import SecretClient
try:
    from azure.storage.blob import BlobServiceClient
except ImportError:
    BlobServiceClient = None


WORD_LIMIT = 50000
MAX_MESSAGES = 20  # <--- Limit the number of messages in memory

# **New Imports for SocketIO**
from flask_socketio import SocketIO, emit, join_room  # <-- Added


class ChatCog:
    def __init__(self, app_instance, flask_app, socketio):  # <-- Modified to accept socketio
        self.bp = Blueprint("chat_blueprint", __name__)
        self.socketio = socketio  # <-- Store socketio instance

        # Initialize OpenAI client
        openai_key = os.getenv("OPENAI_KEY")
        google_key = os.getenv("GOOGLE_API_KEY")
        openai.api_key = openai_key  # Use the fetched secret directly
        self.client = openai

        # Initialize other cogs
        self.web_search_cog = WebSearchCog(openai_client=self.client)
        self.code_files_cog = CodeFilesCog()
        self.orchestration_analysis_cog = OrchestrationAnalysisCog(self.client)

        self.google_key = google_key
        self.app_instance = app_instance

        # self.lang_graph = LangGraph()
        # # Define your tasks
        # self.lang_graph.add_node("handle_user_input")
        # self.lang_graph.add_node("generate_response")
        # self.lang_graph.add_node("process_file_upload")

        # # Define task dependencies
        # self.lang_graph.add_edge("handle_user_input", "generate_response")
        # self.lang_graph.add_edge("process_file_upload", "generate_response")

        # Ensure 'uploads' directory exists
        self.upload_folder = os.path.join(flask_app.instance_path, 'uploads')
        os.makedirs(self.upload_folder, exist_ok=True)
        print(f"Uploads directory set at: {self.upload_folder}")

        # Add Socket.IO event handlers
        self.add_socketio_events()

        self.code_structure_visualizer_cog = CodeStructureVisualizerCog(self.upload_folder)

        # ------------------------------------------------------------
        # Attempt to initialize Azure Storage (optional)
        # If no connection string is found, we keep local storage only
        # ------------------------------------------------------------
        self.use_azure = False
        azure_conn_str = None
        try:
            azure_conn_str = os.getenv("AZURE_BLOB_CONNECTION_STRING")  # Example env var
        except Exception as e:
            print(f'No AZURE_BLOB_CONNECTION_STRING in env', flush=True)

        print(f'azure_conn_str: {azure_conn_str}', flush=True)
        if BlobServiceClient and azure_conn_str:
            try:
                self.blob_service_client = BlobServiceClient.from_connection_string(azure_conn_str)
                self.use_azure = False
                print("Azure Blob Storage is configured. Files will be uploaded to Azure.")
            except Exception as ex:
                print(f"Failed to initialize BlobServiceClient. Using local storage. Error: {ex}")
                self.blob_service_client = False
        else:
            self.blob_service_client = False
            print("Azure Blob Storage is NOT configured. Using local storage by default.")

        self.azure_container_name = os.getenv("AZURE_CONTAINER_NAME", "my-container-name")

        print(f'Adding routes...', flush=True)
        self.add_routes()
        print("Routes added.", flush=True)

    def add_socketio_events(self):
        # @copy_current_request_context
        @self.socketio.on('connect')
        def handle_connect():
            try:
                # Retrieve session_id from WebSocket query parameters
                # session_id = session.get('_id')

                session_id = request.args.get('session_id')
                print(f'session_id connecting: {session_id}', flush=True)

                if not session_id:
                    session_id = str(uuid.uuid4())  # Generate if not provided
                    print(f'Generated new session_id: {session_id}', flush=True)

                join_room(session_id)  # Join WebSocket room with the session_id
                emit('connected', {'session_id': session_id})  # Emit back to client
                print(f"Client connected and joined room: {session_id}", flush=True)


            except Exception as e:
                print(f"Error in connect handler: {e}", flush=True)
                traceback.print_exc()

        @self.socketio.on('verify_room')
        def handle_verify_room(data):
            room = data['room']
            if room in rooms():
                print(f"Client is in room: {room}", flush=True)
            else:
                print(f"Client is NOT in room: {room}", flush=True)

        @self.socketio.on('disconnect')
        def handle_disconnect():
            session_id = request.args.get('session_id')
            print(f'Client disconnected: {session_id}', flush=True)

        @self.socketio.on('orchestrate')
        def handle_orchestration(data):
            session_id = data.get('room')  # Explicitly get session_id from payload
            action = data.get('action')
            if session_id and action:
                print(f'Orchestration event received: {action}', flush=True)
                # Example: Emit a status update for the received action
                status_message = f'Processing action: {action}'
                self.socketio.emit('status_update', {'message': status_message}, room=session_id)
                # Further handling based on action...
            else:
                print('Orchestration event missing session_id or action.', flush=True)


    def _chat_logic(self):
        print(flush=True)
        try:
            # Check the Content-Type of the request
            if request.content_type.startswith('multipart/form-data'):
                # Handle file uploads
                message = request.form.get("message", "")
                model = request.form.get("model", "gpt-4o-mini")
                temperature = float(request.form.get("temperature", 0.7))
                file = request.files.get("file")

                if not file and not message:
                    raise ValueError("No message or file provided")

                session_id = request.form.get("room")
                if not session_id:
                    session_id = str(uuid.uuid4())  # Generate a new session ID if missing
                    print(f"Generated new session_id for /chat: {session_id}", flush=True)

                print(f"Handling multipart/form-data request with session_id: {session_id}", flush=True)

            elif request.is_json:
                # Handle JSON requests
                data = request.get_json()
                if not data:
                    raise ValueError("Invalid JSON payload")

                session_id = data.get("room")
                message = data.get("message", "")
                model = data.get("model", "gpt-4o-mini")
                temperature = float(data.get("temperature", 0.7))
                file = None  # No file in JSON requests

                if not session_id:
                    session_id = str(uuid.uuid4())  # Generate a new session ID if missing
                    print(f"Generated new session_id for /chat: {session_id}", flush=True)

                print(f"Handling application/json request with session_id: {session_id}", flush=True)

            else:
                raise ValueError("Unsupported Content-Type")

            # Retrieve system prompt
            system_prompt = self.get_system_prompt()
            # print(f"System Prompt: {system_prompt}", flush=True)

            # Retrieve other parameters
            message, model, temperature, file = self.get_request_parameters()
            # print(f"Model: {model}, Temperature: {temperature}", flush=True)
            print(f"User Message: {message}", flush=True)

            # ---------------------------------------------------------------
            # Use updated process_uploaded_file that conditionally uses Azure
            # ---------------------------------------------------------------

            file_content, file_url, file_type, uploaded_file = self.process_uploaded_file(
                file=file,
                upload_folder=self.upload_folder,    # local fallback folder
                session_id=session_id,
                # db=db,
                # Toggle Azure settings
                use_azure=self.use_azure,            # only True if we found a conn_str
                blob_service_client=self.blob_service_client,
                container_name=self.azure_container_name
            )

            print('File upload portion cleared.', flush=True)
            if not message and not file_url:
                return jsonify({"error": "No message or file provided"}), 400

            # Manage conversation
            print('Getting conversatoin ID.', flush=True)
            try:
                conversation_id, conversation = self.manage_conversation(session_id)
                if not conversation:
                    raise ValueError("Invalid or missing conversation.")
            except Exception as e:
                print(f"Exception in manage_conversation: {e}", flush=True)
                traceback.print_exc()
                # optionally re-raise
                raise
            if not conversation:
                return jsonify({"error": "Conversation not found or unauthorized"}), 404

            # Get conversation history and truncate if needed
            print('Getting conversation history.', flush=True)
            conversation_history = self.get_conversation_history(conversation_id)
            if len(conversation_history) > MAX_MESSAGES:
                conversation_history = conversation_history[-MAX_MESSAGES:]

            # Analyze user orchestration
            print('Sending for orchestration.', flush=True)
            orchestration = self.orchestration_analysis_cog.analyze_user_orchestration(
                user_message=message,
                conversation_history=conversation_history,
                session_id=session_id
            )

            if not message:
                message = f'User is uploading a file. Respond in acknowledgement that a file was uploaded. Here is the file name: {uploaded_file.original_filename}'

            # print(f"Orchestration: {orchestration}", flush=True)

            # Determine status message based on orchestration
            if orchestration.get("internet_search"):
                status_message = "Searching the internet..."
            elif orchestration.get("image_generation"):
                status_message = "Creating the image..."
            elif orchestration.get("code_intent"):
                status_message = "Processing your code request..."
            elif orchestration.get("file_orchestration"):
                status_message = "Analyzing the uploaded file..."
            else:
                status_message = "Assistant is thinking..."

            # Emit status update via SocketIO
            print(f'status_message: {status_message}', flush=True)
            print(f'Emitting with session id: {session_id}', flush=True)
            self.socketio.emit('status_update', {'message': status_message}, room=session_id)

            # Handle orchestration-specific actions
            if orchestration.get("image_generation", False):
                response = self.handle_image_generation(orchestration, message, conversation_history, conversation_id)
                print(f'response: {response}', flush=True)
                # Emit task completion
                self.socketio.emit('task_complete', {'answer': response['assistant_reply']}, room=session_id)
                return response
            elif orchestration.get("code_structure_orchestration", False):
                response = self.handle_code_structure_visualization(orchestration, message, conversation_history, conversation_id)
                # Emit task completion
                self.socketio.emit('task_complete', {'answer': response['assistant_reply']}, room=session_id)
                return response
            else:
                # Handle other orchestrations
                supplemental_information, assistant_reply = self.handle_orchestration(orchestration, session_id, conversation_id)

                # Prepare messages for OpenAI API
                messages = self.prepare_messages(system_prompt, conversation_history, supplemental_information, message)

                # Trim conversation if necessary (token-based)
                messages = self.trim_conversation(messages, WORD_LIMIT)

                # Generate chat response
                assistant_reply = generate_chat_response(self.client, messages, model, temperature)
                # print(f"Assistant Reply: {assistant_reply}", flush=True)

                # Save messages to the database
                self.save_messages(conversation_id, "user", message)
                self.save_messages(conversation_id, "assistant", assistant_reply)

                # Emit task completion via SocketIO
                self.socketio.emit('task_complete', {'answer': assistant_reply}, room=session_id)

                del messages
                gc.collect()

                return jsonify({
                    "user_message": message,
                    "assistant_reply": assistant_reply,
                    "conversation_history": conversation_history,  # truncated in memory above
                    "orchestration": orchestration,
                    "fileUrl": uploaded_file.file_url if uploaded_file else None,
                    "fileName": uploaded_file.original_filename if uploaded_file else None,
                    "fileType": uploaded_file.file_type if uploaded_file else None,
                    "fileId": uploaded_file.id if uploaded_file else None
                }), 200

        except Exception as e:
            print(f"Error in /chat route: {e}", flush=True)
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500


    def add_routes(self):
        @self.bp.route("/chat", methods=["POST"])
        def chat():
            return self._chat_logic()


        # **New Route to Serve Uploaded Images (or any file) Locally**
        @self.bp.route('/uploads/<filename>')
        def uploaded_file(filename):
            """
            Serve uploaded files from the uploads directory (local).
            If using Azure, you'd serve via a blob URL. This is for local fallback.
            """
            try:
                filename = secure_filename(filename)
                return send_from_directory(self.upload_folder, filename)
            except Exception as e:
                print(f"Error serving file {filename}: {e}", flush=True)
                return jsonify({"error": "File not found."}), 404


        # ############################################################################
        # # NEW Conversation routes go below
        # ############################################################################
        @self.bp.route("/conversations", methods=["GET"])
        def get_conversations():
            try:
                session_id = request.args.get("session_id")  # Obtain from request or context
                print(f"Fetching conversation for session_id: {session_id}", flush=True)
                
                if not session_id:
                    return jsonify({"error": "Session ID is required"}), 400

                print('Is this where the error is coming from???', flush=True)
                # Manage conversation
                print(f'session_id: {session_id}', flush=True)
                conversation_id, conversation_obj = self.manage_conversation(session_id)
                if not conversation_obj:
                    print("error\": \"Conversation not found or unauthorized", flush=True)
                    return jsonify({"error": "Conversation not found or unauthorized"}), 404

                print('Getting conversation history.', flush=True)
                conversation = self.get_conversation_history(conversation_id)
                
                if not conversation:
                    print(f"No conversation found for session_id: {session_id}", flush=True)
                    return jsonify({"conversations": []}), 200
                
                data = {
                    "id": conversation_obj.id,
                    "session_id": conversation_obj.session_id,
                    "title": conversation_obj.title,
                    "timestamp": conversation_obj.timestamp.isoformat() if conversation_obj.timestamp else None,
                    "conversation_history": conversation  # NEW: Include full history
                }
                    
                return jsonify({"conversation": data}), 200

            except Exception as e:
                print(f"Error retrieving conversation: {e}", flush=True)
                traceback.print_exc()
                return jsonify({"error": "Failed to retrieve conversation"}), 500



        @self.bp.route("/conversations/new", methods=["POST"])
        def create_new_conversation():
            """
            Creates a new conversation row with a new session_id, 
            then returns the new conversation data (including session_id).
            """
            try:
                data = request.get_json()
                title = data.get("title", "New Conversation")

                # Generate a brand-new session_id so it doesn't share history
                import uuid
                new_session_id = str(uuid.uuid4())

                # Use your existing get_or_create_conversation method,
                # or just create a new row directly:
                new_conversation = self.get_or_create_conversation(new_session_id, title=title)

                # Return the newly created conversation
                return jsonify({
                    "id": new_conversation.id,
                    "session_id": new_conversation.session_id,
                    "title": new_conversation.title
                }), 200

            except Exception as e:
                print(f"Error creating conversation: {e}", flush=True)
                return jsonify({"error": str(e)}), 500

        @self.bp.route("/ping", methods=["GET"])
        def ping():
            try:
                db.session.execute("SELECT 1")  # quick check
                return "OK", 200
            except:
                db.session.rollback()
                return "DB Connection Lost, Reset", 500




    # Additional Helper Methods

    def get_system_prompt(self):
        if request.content_type.startswith('multipart/form-data'):
            return request.form.get("system_prompt", "You are a USMC AI agent. Provide relevant responses.")
        elif request.is_json:
            try:
                data = request.get_json()
                if not data:
                    return jsonify({"error": "Invalid JSON payload"}), 400
            except Exception as e:
                return jsonify({"error": f"Malformed JSON: {str(e)}"}), 400
            return data.get("system_prompt", "You are a USMC AI agent. Provide relevant responses.")
        else:
            return "You are a USMC AI agent. Provide relevant responses."

    def get_request_parameters(self):
        if request.content_type.startswith('multipart/form-data'):
            message = request.form.get("message", "")
            model = request.form.get("model", "gpt-4o-mini")
            try:
                temperature = float(request.form.get("temperature", 0.7))
            except ValueError:
                temperature = 0.7
            file = request.files.get("file", None)
        elif request.is_json:
            try:
                data = request.get_json()
                if not data:
                    return jsonify({"error": "Invalid JSON payload"}), 400
            except Exception as e:
                return jsonify({"error": f"Malformed JSON: {str(e)}"}), 400
            message = data.get("message", "")
            model = data.get("model", "gpt-4o-mini")
            try:
                temperature = float(data.get("temperature", 0.7))
            except ValueError:
                temperature = 0.7
            file = None
        else:
            message = ""
            model = "gpt-4o-mini"
            temperature = 0.7
            file = None
        return message, model, temperature, file

    def get_or_create_conversation(self, session_id, title="New Conversation", limit=100):
        """
        Fetch an existing conversation by session_id or create a new one.
        """
        try:
            print(f"[get_or_create_conversation] Querying for session_id: {session_id}")
            
            if db.session.query(Conversation).filter_by(session_id=session_id).first() is not None:
                conversation = Conversation.query.filter_by(session_id=session_id).execution_options(timeout=5).first()
            else:
                conversation = None
            
            if conversation:
                print(f"[get_or_create_conversation] Found conversation ID: {conversation.id}")
            else:
                print(f"[get_or_create_conversation] No conversation found for session_id: {session_id}. Creating a new one.")
                conversation = Conversation(session_id=session_id, title=title)
                db.session.add(conversation)
                print(f"[get_or_create_conversation] Added new conversation to session.")
                db.session.commit()  # Save changes to the database
                # print(f"[get_or_create_conversation] Committed new conversation with ID: {conversation.id}")
        
            return conversation
        
        except Exception as e:
            # Handle potential database issues
            db.session.rollback()  # Rollback the transaction in case of failure
            print(f"[get_or_create_conversation] Failed to get or create conversation: {e}")
            traceback.print_exc()
            raise


    
    def manage_conversation(self, session_id, limit=100):
        """
        Fetches or creates a conversation based on session_id.
        Returns a tuple of (conversation_id, Conversation object).
        """
        try:
            # Debugging: print all conversations (potentially large if the table is big)
            # all_conversations = Conversation.query.all()
            # for conversation in all_conversations:
            #     print(
            #         f"ID: {conversation.id}, "
            #         f"Session ID: {conversation.session_id}, "
            #         f"Title: {conversation.title}, "
            #         f"Timestamp: {conversation.timestamp}",
            #         flush=True
            #     )

            print(f"[manage_conversation] Managing conversation for session_id: {session_id}")

            if db.session.query(Conversation).filter_by(session_id=session_id).first() is not None:
                conversation = db.session.query(Conversation).filter_by(session_id=session_id).execution_options(timeout=5).first()
                print(f"[manage_conversation] conversation found => ID: {conversation.id if conversation else None}", flush=True)
            else:
                conversation = None


            # If no conversation is found, create a new one
            if not conversation:
                print(f"[manage_conversation] No conversation found for session_id: {session_id}. Creating a new one.", flush=True)
                conversation = Conversation(session_id=session_id, title="New Conversation")
                db.session.add(conversation)
                db.session.commit()  # commit the create
                print(f"[manage_conversation] Created new conversation with ID: {conversation.id}", flush=True)
            else:
                # Optionally refresh the conversation
                db.session.refresh(conversation)
                print(f"[manage_conversation] Retrieved conversation ID: {conversation.id}", flush=True)

            return conversation.id, conversation

        except Exception as e:
            db.session.rollback()  # Make sure to rollback on error
            print(f"[manage_conversation] Error: {e}")
            traceback.print_exc()
            return None, None



    def get_conversation_history(self, conversation_id, limit=50, offset=0):
        """
        Retrieve messages from the database with pagination and return them as a list of {role, content}.
        """
        try:
            if db.session.query(Message).filter_by(conversation_id=conversation_id).first() is not None:
                messages_db = (
                    Message.query
                    .filter_by(conversation_id=conversation_id)
                    .order_by(Message.timestamp)
                    .offset(offset)
                    .limit(limit)
                    .all()
                )
            else:
                print(f"Error fetching conversation history: {e}", flush=True)
                return []
            return [{"role": msg.role, "content": msg.content} for msg in messages_db]
        except Exception as e:
            print(f"Error fetching conversation history: {e}", flush=True)
            return []



    def handle_orchestration(self, orchestration, session_id=None, conversation_id=None):
        supplemental_information = {}
        assistant_reply = ""
        
        if orchestration.get("file_orchestration", False):
            supplemental_information, assistant_reply = self.handle_file_orchestration(orchestration, session_id)
        elif orchestration.get("code_orchestration", False):
            code_content = self.code_files_cog.get_all_code_files_content()
            if code_content:
                supplemental_information = {
                    "role": "system",
                    "content": (
                        f"\n\nYou have been supplemented with information from your code base to answer this query.\n***{code_content}***"
                    )
                }
            else:
                assistant_reply = "No code files found to provide."
        elif orchestration.get("internet_search", False):
            query = request.json.get("message", "")
            search_content = self.web_search_cog.web_search(query, self.get_conversation_history(conversation_id))
            sys_search_content = (
                'Do not say "I am unable to browse the internet," because you have information directly retrieved from the internet. '
                'Give a confident answer based on the suplimented information. Only use the most relevant and accurate information that matches the User Query. '
                'Always include the source with the provided url as [source](url) at the end of each section you used to generate the response.'
            )
            print(flush=True)
            print(f'search_content:\n\n{search_content}')
            print(flush=True)
            supplemental_information = {
                "role": "system",
                "content": (
                    f"{sys_search_content}\n\nInternet Content:\n***{search_content}***"
                )
            }
        elif orchestration.get("rand_num", []):
            numbers = orchestration["rand_num"]
            if len(numbers) == 2:
                import random
                rand_num = random.randint(numbers[0], numbers[1])
                assistant_reply = f"Your random number between {numbers[0]} and {numbers[1]} is {rand_num}."
            else:
                assistant_reply = "Please provide a valid range for the random number."
        return supplemental_information, assistant_reply
        
    def handle_file_orchestration(self, orchestration, session_id):
        """
        Handle file orchestration based on the orchestration instructions.

        Args:
            orchestration (dict): The orchestration JSON object containing directives.
            session_id (str): The current session ID.

        Returns:
            tuple: A tuple containing supplemental information (dict) and assistant reply (str).
        """
        supplemental_information = {}
        assistant_reply = ""
        file_ids = orchestration.get("file_ids", [])

        print("Starting handle_file_orchestration", flush=True)

        try:
            # Fetch all uploaded files for the current session
            print("Fetching uploaded files for session_id:", session_id, flush=True)
            uploaded_files = UploadedFile.query.filter_by(session_id=session_id).all()
            uploaded_file_ids = {str(file.id): file for file in uploaded_files}
            print(f'uploaded_file_ids: {uploaded_file_ids}', flush=True)

            # If file orchestration is not requested, return empty responses
            if not orchestration.get("file_orchestration", False):
                print("File orchestration not requested.", flush=True)
                return supplemental_information, assistant_reply  # No file orchestration needed

            # Case 1: General file query (no specific file_ids provided)
            if not file_ids:
                print("Handling general file query (no specific file_ids provided).", flush=True)
                if not uploaded_files:
                    print("No uploaded files found.", flush=True)
                    assistant_reply = "No files have been uploaded yet."
                    return supplemental_information, assistant_reply
                else:
                    try:
                        file_list_str = "\n".join([f"- {file.original_filename} (ID: {file.id})" for file in uploaded_files])
                        print("Constructed file_list_str for general query:", file_list_str, flush=True)
                        assistant_reply = f"Here are the uploaded file names:\n{file_list_str}"
                        supplemental_information = {
                            "role": "system",
                            "content": (
                                '\n\nYou are being supplemented with the following information.\n'
                                f"List of uploaded file names:\n***{file_list_str}***"
                            )
                        }
                        print("General file query response constructed successfully.", flush=True)
                    except Exception as e:
                        print(f"Error constructing general file list: {e}", flush=True)
                        assistant_reply = "An error occurred while listing the uploaded files."
                    return supplemental_information, assistant_reply

            # Case 2: Specific file query
            print("Handling specific file query with file_ids:", file_ids, flush=True)
            # Validate requested file IDs against uploaded files
            valid_requested_file_ids = [fid for fid in file_ids if fid in uploaded_file_ids]
            invalid_file_ids = [fid for fid in file_ids if fid not in uploaded_file_ids]
            print(f"Valid file_ids: {valid_requested_file_ids}", flush=True)
            print(f"Invalid file_ids: {invalid_file_ids}", flush=True)

            if not valid_requested_file_ids:
                # No valid file IDs found
                if invalid_file_ids:
                    print("No valid uploaded files found for the requested file IDs.", flush=True)
                    assistant_reply = f"No valid uploaded files found for the requested file IDs: {', '.join(invalid_file_ids)}."
                else:
                    assistant_reply = "No valid uploaded files found for the requested file IDs."
                return supplemental_information, assistant_reply

            # Determine the number of valid requested files
            num_requested_files = len(valid_requested_file_ids)
            print(f"Number of valid requested files: {num_requested_files}", flush=True)

            if num_requested_files > 3:
                # More than 3 files requested: only list file names
                print("More than 3 files requested. Listing file names without contents.", flush=True)
                try:
                    file_list_str = "\n".join([f"- {uploaded_file_ids[fid].original_filename} (ID: {fid})" for fid in valid_requested_file_ids])
                    print("Constructed file_list_str for multiple specific files:", file_list_str, flush=True)
                    assistant_reply = f"Here are the requested file names:\n{file_list_str}\n\nNote: File contents are not displayed as more than 3 files were requested."
                    supplemental_information = {
                        "role": "system",
                        "content": (
                            '\n\nYou are being supplemented with the following information.\n'
                            f"List of requested file names:\n***{file_list_str}***"
                        )
                    }
                    # If there are invalid file IDs, append them to the assistant reply
                    if invalid_file_ids:
                        print("Appending information about invalid file_ids.", flush=True)
                        assistant_reply += f"\nAdditionally, the following requested file IDs are invalid: {', '.join(invalid_file_ids)}."
                except Exception as e:
                    print(f"Error constructing file list for multiple specific files: {e}", flush=True)
                    assistant_reply = "An error occurred while listing the requested files."
                return supplemental_information, assistant_reply

            # If 1-3 files are requested: include file names and contents
            print("1-3 files requested. Including file names and contents.", flush=True)
            file_contents = []
            errors = []

            for fid in valid_requested_file_ids:
                uploaded_file = uploaded_file_ids.get(fid)
                if uploaded_file:
                    file_path = os.path.join(self.upload_folder, uploaded_file.filename)
                    print(f"Processing file: {uploaded_file.original_filename} (ID: {fid}) at path: {file_path}", flush=True)
                    if os.path.exists(file_path):
                        try:
                            print(f"File exists. Processing file: {uploaded_file.original_filename}", flush=True)
                            # Process the uploaded file (adjust parameters as needed)
                            # read=True => means read from local path
                            file_content = self.process_uploaded_file(
                                file=None,
                                upload_folder=self.upload_folder,
                                session_id=uploaded_file.session_id,
                                # db=db,
                                read=True,
                                path=file_path,
                                # Still pass Azure info, but read=True means we won't reupload
                                use_azure=self.use_azure,
                                blob_service_client=self.blob_service_client,
                                container_name=self.azure_container_name
                            )
                            print(f"Successfully processed file: {uploaded_file.original_filename}", flush=True)
                            file_contents.append((uploaded_file.original_filename, file_content))
                        except Exception as e:
                            print(f"Error processing file {uploaded_file.original_filename}: {e}", flush=True)
                            errors.append(f"Error processing file '{uploaded_file.original_filename}'.")
                    else:
                        print(f"File not found on server: {uploaded_file.original_filename}", flush=True)
                        errors.append(f"File '{uploaded_file.original_filename}' not found on server.")
                else:
                    print(f"Uploaded file with ID '{fid}' not found.", flush=True)
                    errors.append(f"Uploaded file with ID '{fid}' not found.")

            # Construct the assistant reply with file contents
            if file_contents:
                print("Constructing assistant reply with file contents.", flush=True)
                try:
                    for fname, content in file_contents:
                        assistant_reply += f"**{fname}:**\n{content}\n\n"
                except Exception as e:
                    print(f"Error constructing assistant reply with file contents: {e}", flush=True)
                    assistant_reply += "An error occurred while constructing the file contents in the reply.\n"

            # Append any errors encountered during processing
            if errors:
                print("Appending errors to assistant reply.", flush=True)
                try:
                    assistant_reply += "\n".join(errors)
                except Exception as e:
                    print(f"Error appending errors to assistant reply: {e}", flush=True)
                    assistant_reply += "\nAn error occurred while appending error messages."

            # Construct the supplemental information with file contents
            if file_contents:
                print("Constructing supplemental information with file contents.", flush=True)
                try:
                    file_content_str = "\n\n".join([
                        f"File: {fname}\nContent:\n***{content}***" for fname, content in file_contents
                    ])
                    supplemental_information = {
                        "role": "system",
                        "content": (
                            '\n\nYou are being supplemented with the following information from the files.\n'
                            f"{file_content_str}"
                        )
                    }
                except Exception as e:
                    print(f"Error constructing supplemental information: {e}", flush=True)
                    # If there's an error, do not include supplemental information

            # If there are invalid file IDs, inform the user
            if invalid_file_ids:
                print("Informing user about invalid file_ids.", flush=True)
                try:
                    assistant_reply += f"\nAdditionally, the following requested file IDs are invalid: {', '.join(invalid_file_ids)}."
                except Exception as e:
                    print(f"Error appending invalid file IDs to assistant reply: {e}", flush=True)
                    assistant_reply += "\nAn error occurred while informing about invalid file IDs."

            print("handle_file_orchestration completed successfully.", flush=True)
            return supplemental_information, assistant_reply
        except Exception as e:
            print('Error:', e)
            return supplemental_information, assistant_reply
            

    def handle_image_generation(self, orchestration, user_message, conversation_history, conversation_id):
        """
        Handles image generation and returns a response immediately.
        """
        supplemental_information = {}
        assistant_reply = ""
        prompt = orchestration.get("image_prompt", "")
        if prompt:
            image_url = generate_image(prompt, self.client)
            assistant_reply = f"![Generated Image]({image_url})"
            conversation_history.append({"role": "assistant", "content": assistant_reply})
            self.save_messages(conversation_id, "assistant", assistant_reply)
        else:
            assistant_reply = "No image prompt provided."
            conversation_history.append({"role": "assistant", "content": assistant_reply})
            self.save_messages(conversation_id, "assistant", assistant_reply)
        
        return jsonify({
            "user_message": user_message,
            "assistant_reply": assistant_reply,
            "conversation_history": conversation_history,
            "orchestration": orchestration,
            "fileUrl": None,
            "fileName": None,
            "fileType": None
        })

    def handle_code_structure_visualization(self, orchestration, user_message, conversation_history, conversation_id):
        """
        Handles code structure visualization and returns a response immediately.
        """
        supplemental_information = {}
        assistant_reply = ""
        image_url = self.code_structure_visualizer_cog.generate_codebase_structure_diagram()
        if image_url:
            assistant_reply = f"![Codebase Structure]({image_url})"
            conversation_history.append({"role": "assistant", "content": assistant_reply})
            self.save_messages(conversation_id, "assistant", assistant_reply)
        else:
            assistant_reply = "Failed to generate codebase structure diagram."
            conversation_history.append({"role": "assistant", "content": assistant_reply})
            self.save_messages(conversation_id, "assistant", assistant_reply)
        
        # Optionally add code content
        code_content = self.code_files_cog.get_all_code_files_content()
        if code_content:
            supplemental_information = {
                "role": "system",
                "content": (
                    f"\n\nYou have been supplemented with information from your code base to answer this query.\n***{code_content}***"
                )
            }
        
        return jsonify({
            "user_message": user_message,
            "assistant_reply": assistant_reply,
            "conversation_history": conversation_history,
            "orchestration": orchestration,
            "fileUrl": None,
            "fileName": None,
            "fileType": None
        })

    def prepare_messages(self, system_prompt, conversation_history, supplemental_information, user_message):
        additional_instructions = (
            "Generate responses as structured and easy-to-read.  \n"
            "Provide responses using correct markdown formatting. It is critical that markdown format is used.  \n"
            "Use headings (e.g., ## Section Title), numbered lists, and bullet points to format output when the response is more than 2 sentences long. If 1-3 sentence, do not use a heading!  \n"
            "Ensure sufficient line breaks between sections to improve readability. Generally, limit responses to no more than 1500 tokens."
        )
        messages = [
            {
                "role": "system", 
                "content": f"Your role is:\n{system_prompt} \n\nStructured response Guidelines:\n{additional_instructions}"
            }
        ] + conversation_history

        if supplemental_information:
            messages.append(supplemental_information)

        messages.append({"role": "user", "content": user_message})
        return messages

    def trim_conversation(self, messages, max_tokens=WORD_LIMIT):
        """
        Trim the conversation by token count if it exceeds WORD_LIMIT.
        """
        
        encoding = tiktoken.encoding_for_model("gpt-4o-mini")
        total_tokens = 0
        trimmed = []
        
        for message in reversed(messages):
            message_tokens = len(encoding.encode(json.dumps(message)))
            if total_tokens + message_tokens > max_tokens:
                break
            trimmed.insert(0, message)
            total_tokens += message_tokens
        
        if not trimmed and messages:
            trimmed.append(messages[-1])
        
        return trimmed

    def save_messages(self, conversation_id, role, content):
        """Save a message to the database."""
        msg = Message(
            conversation_id=conversation_id,
            role=role,
            content=content
        )
        db.session.add(msg)
        db.session.commit()



    def process_uploaded_file(
        self,
        file=None,
        upload_folder=None,
        session_id=None,
        # db=None,
        read=False,
        path=None,
        use_azure=False,
        blob_service_client=None,
        container_name=None
    ):
        """
        Handles file saving and processing.
        
        By default, files are stored locally (in `upload_folder`). 
        If `use_azure=True` and a valid `blob_service_client` is provided,
        files will be uploaded to Azure Blob Storage instead.

        :param file: File object from the request.
        :param upload_folder: Directory to save local uploads (if not using Azure).
        :param session_id: Current session ID.
        :param db_session: Database session.
        :param read: If True, read the file content (from `path`) but do not upload.
        :param path: Path to the file if reading.
        :param use_azure: Whether to upload to Azure instead of saving locally.
        :param blob_service_client: An instance of BlobServiceClient (if use_azure=True).
        :param container_name: The name of the Azure container (if use_azure=True).
        :return: Tuple (file_content, file_url, file_type, uploaded_file)
        """

        # 1) If 'read=True' and 'path' is provided, we just read the file content from disk
        #    and do not upload anywhere.
        print('In processing file', flush=True)
        if read and path:
            return self.read_file_content(path)

        # 2) If there's no file provided, we can't do anything
        print('No file', flush=True)
        if not file:
            print('Returning, no file', flush=True)
            return '', None, None, None

        print('Not returning, continuing to process files.', flush=True)
        # Generate secure filename
        filename = secure_filename(file.filename)
        unique_filename = f"{uuid.uuid4()}_{filename}"

        # --------------------------
        # A) Upload to Azure Storage
        # --------------------------
        if use_azure and blob_service_client and container_name:
            # Read file content into memory
            file_bytes = file.read()

            # Attempt uploading to Azure
            blob_client = blob_service_client.get_blob_client(
                container=container_name,
                blob=unique_filename
            )
            try:
                blob_client.upload_blob(file_bytes, overwrite=True)
                azure_blob_url = blob_client.url
                file_url = azure_blob_url
            except Exception as e:
                print("Error uploading to Azure Blob Storage:", e)
                return "Error uploading file.", None, None, None
            try:
                # Insert a record in the database with the Azure file info
                uploaded_file = UploadedFile(
                    session_id=session_id,
                    filename=unique_filename,
                    original_filename=filename,
                    file_url=file_url,
                    file_type=file.content_type
                )
                db.session.add(uploaded_file)
                db.session.commit()
            except Exception as e:
                # Rollback the session to avoid leaving it in a broken state
                db.session.rollback()
                print(f"Error while saving uploaded file record to the database: {e}", flush=True)
                traceback.print_exc()
                # Optionally, re-raise the exception to propagate it
                raise e
                


            # Extract the text for indexing/LLM usage
            file_content = self.extract_content_from_memory(
                file_bytes=file_bytes,
                content_type=file.content_type
            )
            file_type = file.content_type

        # -------------------------
        # B) Store File Locally
        # -------------------------
        else:
            # Construct local file path and save the file
            file_path = os.path.join(upload_folder, unique_filename)
            file.save(file_path)
            print(flush=True)

            # Log file save details
            try:
                file_size = os.path.getsize(file_path)
                print(f"Saved file at {file_path}, size: {file_size} bytes", flush=True)
            except Exception as e:
                print(f"Error accessing saved file: {e}", flush=True)
                return "Error accessing saved file.", None, None, None


            print(f'In process files util: {session_id}', flush=True)
            print(f'In process files util. file_path: {file_path}', flush=True)

            print(flush=True)

            # Insert record in DB with local file path
            try:
                uploaded_file = UploadedFile(
                    session_id=session_id,
                    filename=unique_filename,
                    original_filename=filename,
                    file_url=f"/uploads/{unique_filename}",
                    file_type=file.content_type
                )
                print(f'Adding to db: {uploaded_file}', flush=True)
                db.session.add(uploaded_file)
                db.session.commit()
            except Exception as e:
                # Rollback the session to avoid leaving it in a broken state
                db.session.rollback()
                print(f"Error while saving uploaded file record to the database: {e}", flush=True)
                traceback.print_exc()
                # Optionally, re-raise the exception to propagate it
                raise e

            print(f'Added to db: {uploaded_file}', flush=True)

            # Extract text from local file
            if file.content_type == 'application/pdf':
                file_content = self.extract_text_from_pdf(file_path)
            elif file.content_type in [
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/msword'
            ]:
                file_content = self.extract_text_from_docx(file_path)
            elif file.content_type in [
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                'application/vnd.ms-excel'
            ]:
                file_content = extract_text_from_excel(file_path)
            else:
                # Attempt reading as text
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        file_content = f.read()
                except Exception as e:
                    print("Error reading file:", e)
                    file_content = "Error processing file."

            # Local URL for the stored file
            file_url = f"/uploads/{unique_filename}"
            file_type = file.content_type

        return file_content, file_url, file_type, uploaded_file


    def read_file_content(self, path):
        """
        Reads an existing file from a local path, returning its text content.
        If you want to read from Azure for an existing file, you'll need to
        download to a temp file or memory first (not implemented here).
        """
        file_extension = os.path.splitext(path)[1].lower()
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(path)
        elif file_extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(path)
        elif file_extension in ['.xlsx', '.xls']:
            return self.extract_text_from_excel(path)
        else:
            try:
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            except Exception as e:
                print("Error reading file:", e)
                return "Error processing file."


    # -----------------------
    # In-memory extractions
    # -----------------------
    def extract_content_from_memory(self, file_bytes, content_type):
        """
        Reads the file content from an in-memory bytes object.
        """
        if not file_bytes:
            return ""

        if content_type == 'application/pdf':
            return self.extract_pdf_from_memory(file_bytes)
        elif content_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]:
            return self.extract_docx_from_memory(file_bytes)
        elif content_type in [
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel'
        ]:
            return self.extract_excel_from_memory(file_bytes)
        else:
            # Assume it's plain text
            try:
                return file_bytes.decode('utf-8', errors='ignore')
            except Exception as e:
                print("Error decoding in-memory text file:", e)
                return "Error processing file."


    def extract_pdf_from_memory(self, file_bytes):
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            file_content = ""
            for page in reader.pages:
                file_content += page.extract_text() or ""
            return self.truncate_content(file_content)
        except Exception as e:
            print("Error reading PDF from memory:", e)
            return "Error processing PDF file."


    def extract_docx_from_memory(self, file_bytes):
        try:
            doc = Document(io.BytesIO(file_bytes))
            file_content = "\n".join([p.text for p in doc.paragraphs])
            return self.truncate_content(file_content)
        except Exception as e:
            print("Error reading DOCX from memory:", e)
            return "Error processing Word file."


    def extract_excel_from_memory(self, file_bytes):
        try:
            wb = load_workbook(filename=io.BytesIO(file_bytes))
            sheet = wb.active
            file_content = ""
            for row in sheet.iter_rows(values_only=True):
                file_content += ' '.join(str(cell) for cell in row if cell is not None) + "\n"
            return self.truncate_content(file_content)
        except Exception as e:
            print("Error reading Excel from memory:", e)
            return "Error processing Excel file."


    # -----------------------
    # Local file extractions
    # -----------------------
    def extract_text_from_pdf(self, file_path):
        try:
            reader = PdfReader(file_path)
            file_content = ""
            for page in reader.pages:
                file_content += page.extract_text() or ""
            return self.truncate_content(file_content)
        except Exception as e:
            print("Error reading PDF:", e)
            return "Error processing PDF file."


    def extract_text_from_docx(self, file_path):
        try:
            doc = Document(file_path)
            file_content = "\n".join([p.text for p in doc.paragraphs])
            return self.truncate_content(file_content)
        except Exception as e:
            print("Error reading DOCX:", e)
            return "Error processing Word file."


    def extract_text_from_excel(self, file_path):
        try:
            wb = load_workbook(file_path)
            sheet = wb.active
            file_content = ""
            for row in sheet.iter_rows(values_only=True):
                # Convert each cell to string if not None, then join
                file_content += ' '.join(str(cell) for cell in row if cell is not None) + "\n"
            return self.truncate_content(file_content)
        except Exception as e:
            print("Error reading Excel file:", e)
            return "Error processing Excel file."


    def truncate_content(self, content):
        """
        Utility to ensure we don't exceed the WORD_LIMIT.
        """
        words = content.split()
        if len(words) > WORD_LIMIT:
            return ' '.join(words[:WORD_LIMIT]) + "\n\n[Text truncated after 50,000 words.]"
        return content
